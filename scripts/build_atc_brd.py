#!/usr/bin/env python3
"""Populate Master BRD template with Air Traffic Control Marketing Data Governance content."""

from copy import deepcopy
from docx import Document
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

TEMPLATE = "Master BRD Marketing Operations FY25.docx"
OUTPUT = "BRD - Air Traffic Control Marketing Data Governance FY25.docx"

REPLACEMENTS = {
    "<Include the reason why the new requirements are needed and the impact it is having on the business>": (
        "Marketing activity data is collected across multiple channels, platforms, and ingestion paths "
        "(online forms, manual uploads, and content syndication partners) without a single governed "
        "definition of required data elements, valid business use, or downstream impact. This fragmentation "
        "creates inconsistent attribution, unreliable operational reporting, gaps in lead scoring inputs, "
        "and limited seller enablement context when new campaigns or projects are launched.\n\n"
        "The Air Traffic Control (ATC) initiative addresses this by standardizing how marketing activities "
        "flow through an orchestration layer (Tray.io) into Snowflake, and by maintaining a baseline registry "
        "of data elements with documented business purpose across four critical consumption areas: "
        "Attribution, Operational Reporting, Lead Scoring, and Seller Enablement."
    ),
    "<Include the purpose, goals and audience for the document> ": (
        "This Business Requirements Document (BRD) defines the business requirements for marketing data "
        "governance under the Air Traffic Control program. It establishes the vision, scope, process "
        "changes, stakeholder needs, and success criteria for governing marketing data from capture "
        "through publication in Snowflake.\n\n"
        "Primary audience: Marketing Operations, Marketing Analytics, Data Engineering, Demand Generation, "
        "Sales Operations, Compliance, and program stakeholders responsible for approving new marketing "
        "activities and data changes."
    ),
    "<Include who will benefit from the change and the expected impacts to the company>": (
        "Marketing Operations will benefit from a single source of truth for required fields, data rules, "
        "and reporting purpose by activity type.\n\n"
        "Attribution and Analytics teams will benefit from consistent tracking parameters and validated "
        "value sources, improving campaign and channel performance measurement.\n\n"
        "Lead Scoring and Sales teams will benefit from reliable contact, account, and activity detail "
        "fields that support qualification and seller enablement.\n\n"
        "Compliance and Privacy stakeholders will benefit from documented consent, permissions, and data "
        "classification aligned to each use case.\n\n"
        "Expected company impact: faster onboarding of new marketing programs, reduced rework from "
        "missing or misaligned data, improved reporting confidence, and a repeatable governance process "
        "when new projects affect downstream marketing data consumers."
    ),
    "<Include any high-level assumptions, responsibilities, constraints, milestones and delieverables that are needed to ensure the success of the project>": (
        "In scope for FY25 Air Traffic Control marketing data governance:\n"
        "• Baseline data element registry for marketing activities across Manual Uploads, Online Forms, "
        "and Content Syndication processes.\n"
        "• Use case coverage for Offers (Contact Us, Demo, Trial, Promotion, Other), Events (1st Party, "
        "3rd Party, Sponsorships), Webinars, Paid Leads, and related activity types.\n"
        "• Standardized field sections: System Classification, Contact Identity, Account Identity, "
        "Compliance/Permissions, Legacy and Future Tracking Parameters, Event Tracking, and Activity Detail.\n"
        "• Documentation of value source, data rules, input type, and reporting/process purpose per field.\n"
        "• Mapping of each data element to Attribution, Operational Reporting, Lead Scoring, and Seller "
        "Enablement consumption.\n"
        "• Orchestration of marketing activity data through Tray.io with publication to Snowflake.\n"
        "• Governance workflow for business confirmation, stakeholder review, and change control when "
        "new projects are introduced."
    ),
    "<Anything that does not fall within the established scope of a project>": (
        "• Implementation of Tray.io workflows beyond marketing activity data ingestion and publication.\n"
        "• Full Snowflake data model redesign or enterprise-wide master data management outside marketing.\n"
        "• Replacement of source systems (e.g., MAP, CRM, event platforms) or vendor contract changes.\n"
        "• Historical backfill of legacy marketing data unless explicitly approved as a separate workstream.\n"
        "• Non-marketing data domains (finance, product telemetry, HR, etc.).\n"
        "• Final production deployment sign-off for individual channel tools (email, paid social, etc.) "
        "outside the governed data baseline and orchestration path."
    ),
    "<A thing that is accepted as true or as certain to happen, without proof>": (
        "• Tray.io will serve as the orchestration platform for marketing activity data flows.\n"
        "• Snowflake will be the authoritative analytical landing zone for governed marketing activity data.\n"
        "• Business stakeholders will participate in field-level review and confirmation cycles.\n"
        "• Existing source forms and partner feeds can be mapped to the baseline field model.\n"
        "• The baseline registry will be maintained in this repository and updated as new activities launch."
    ),
    "<A situation in which you need something or someone and are unable to continue normally without them>": (
        "• Access to source form exports and partner specifications (e.g., Contact Us Forms, Offer Form "
        "Submissions, Integrate/content syndication inputs).\n"
        "• Marketing Operations ownership of baseline maintenance and change approval.\n"
        "• Data Engineering support for Tray.io-to-Snowflake pipeline configuration and validation.\n"
        "• Analytics and Attribution team input on tracking parameter requirements.\n"
        "• Compliance review for consent, permissions, and PII-related fields.\n"
        "• Platform access for Tray.io and Snowflake in target environments."
    ),
    "<A Key objective is a specific measurable and achievable goal that an individual or team needs to accomplish within a given timeframe>": (
        "1. Establish and maintain a complete baseline of marketing data elements with business purpose "
        "documented for all in-scope activity types.\n"
        "2. Ensure every in-scope field is mapped to at least one downstream use case "
        "(Attribution, Operational Reporting, Lead Scoring, or Seller Enablement) or explicitly flagged "
        "for business confirmation.\n"
        "3. Standardize ingestion through Tray.io orchestration with governed publication to Snowflake.\n"
        "4. Enable impact assessment when new marketing projects or fields are proposed.\n"
        "5. Reduce data quality defects related to missing, mislabeled, or inconsistently sourced fields."
    ),
    "<A Success metric is measurable data used to determine the achievements of your business efforts>": (
        "• 100% of in-scope activity types represented in the baseline data registry.\n"
        "• ≥95% of baseline fields have completed Reporting/Process Purpose documentation.\n"
        "• ≥90% of baseline fields have confirmed Value Source and Data Rules alignment.\n"
        "• Zero net-new production fields introduced without baseline review and downstream impact assessment.\n"
        "• Reduction in attribution and reporting defects tied to missing tracking or identity fields "
        "(baseline vs. post-ATC measurement).\n"
        "• Stakeholder review completion for flagged fields (Business Confirmation Needed = Yes)."
    ),
    "<An As-Is business process is\u00a0a detailed overview of the current state of a company's process, culture, and capabilities. It outlines how work is currently performed and how information flows through the organization>": (
        "Current state (As-Is):\n"
        "1. Marketing activities are launched across multiple channels with channel-specific forms, uploads, "
        "and partner feeds.\n"
        "2. Required data fields vary by activity type and are documented inconsistently across teams and tools.\n"
        "3. Tracking parameters (legacy and future UTM/campaign fields) are not uniformly applied or validated.\n"
        "4. Data flows into downstream systems through disparate paths with limited central governance.\n"
        "5. When new projects launch, downstream impacts to Attribution, Reporting, Lead Scoring, and Seller "
        "Enablement are assessed ad hoc.\n"
        "6. Field purpose, value source, and data rules are maintained in spreadsheets or local documentation "
        "without a single governed repository.\n"
        "7. Business review of field additions (e.g., Contact Us and Offer form fields) occurs reactively "
        "after issues surface in reporting or sales workflows."
    ),
    "<Is\u00a0the future state of an as-is process in the format of a business process flow>": (
        "Future state (To-Be):\n"
        "1. All in-scope marketing activities are onboarded against the ATC baseline field model before launch.\n"
        "2. Activity data is captured at source (forms, uploads, syndication) using standardized sections "
        "and field definitions.\n"
        "3. Tray.io orchestrates validation, transformation, and routing of activity payloads.\n"
        "4. Governed marketing activity data is published to Snowflake for downstream consumption.\n"
        "5. The baseline registry in this repository serves as the authoritative reference for field purpose, "
        "rules, and downstream use case mapping.\n"
        "6. New project requests trigger a structured impact review against Attribution, Operational Reporting, "
        "Lead Scoring, and Seller Enablement.\n"
        "7. Stakeholders complete business confirmation for flagged fields prior to production use.\n\n"
        "To-Be process flow (summary):\n"
        "Source Capture → Tray.io Orchestration → Snowflake Publication → Downstream Consumption "
        "(Attribution / Reporting / Lead Scoring / Seller Enablement)"
    ),
    "<Each requirement for the stakeholder should include an actor, an action, an object and a qualifier if appropriate>": (
        "GR-01: Marketing Operations shall maintain the baseline data registry for all in-scope marketing "
        "activity types and publish updates through the ATC repository.\n"
        "GR-02: Program teams shall assess new marketing projects against the baseline before production launch.\n"
        "GR-03: Data Engineering shall configure Tray.io workflows to ingest, validate, and publish "
        "governed activity data to Snowflake.\n"
        "GR-04: Business stakeholders shall document Reporting/Process Purpose for every field used in "
        "production reporting or operational workflows.\n"
        "GR-05: Compliance shall review fields in Section 4 (Compliance/Permissions) and confirm handling "
        "requirements for consent and marketing permissions.\n"
        "GR-06: Analytics shall validate that Attribution and Operational Reporting flags align with "
        "actual dashboard and model consumption.\n"
        "GR-07: Sales Operations shall validate Seller Enablement field requirements for hand-raiser and "
        "non-hand-raiser use cases.\n"
        "GR-08: The program shall flag fields requiring business confirmation and track resolution status "
        "in the baseline (Business Confirmation Needed, Amanda Reviewed, stakeholder comment columns)."
    ),
}

ONLINE_FORMS_REQUIREMENTS = (
        "Online Forms Requirements (Offers, Contact Us, Webinars, Events):\n"
        "OF-01: Form owners shall map all published form fields to baseline sections (1–8) before go-live.\n"
        "OF-02: Contact and Account Identity fields (Sections 2–3) shall be captured for all hand-raiser "
        "use cases unless explicitly exempted with business approval.\n"
        "OF-03: Future tracking parameters (Section 6) shall be populated per campaign standards for "
        "Attribution and Operational Reporting.\n"
        "OF-04: Activity Detail fields (Section 8) shall document offer-specific, event-specific, or "
        "webinar-specific context required for Lead Scoring and Seller Enablement.\n"
        "OF-05: New form fields identified in source exports (e.g., Offer Form Submissions) shall be added "
        "to the baseline with value source, data rules, and downstream use case flags.\n\n"
        "Manual Uploads and Content Syndication Requirements:\n"
        "MU-01: Manual upload templates shall conform to baseline field names and input types.\n"
        "MU-02: Content syndication partner fields shall map to Integrate Input/Defined or Tray.io/MAT "
        "value sources as documented in the baseline.\n"
        "MU-03: Partner-provided fields without baseline mapping shall be flagged for business confirmation "
        "before inclusion in Snowflake publication.\n"
        "MU-04: Paid lead and syndicated content activities shall include compliance fields where required "
        "for permission-based outreach."
)

# Second occurrence of stakeholder requirements placeholder gets different content
PAID_SOCIAL_REQUIREMENTS = (
    "Attribution and Operational Reporting Requirements:\n"
    "AR-01: Attribution consumers shall receive consistent legacy (Section 5) and future (Section 6) "
    "tracking parameters across all in-scope activity types.\n"
    "AR-02: Operational Reporting shall use documented Reporting/Process Purpose to validate dashboard "
    "and operational metric definitions.\n"
    "AR-03: Fields flagged for Attribution or Operational Reporting shall not be deprecated without "
    "downstream impact assessment and stakeholder sign-off.\n\n"
    "Lead Scoring and Seller Enablement Requirements:\n"
    "LS-01: Lead Scoring shall consume standardized contact, account, funnel stage, and activity detail "
    "fields from governed Snowflake tables.\n"
    "LS-02: Seller Enablement shall receive activity context (e.g., sales conversation interest, discussion "
    "description, thank-you landing page) for offer and event use cases where flagged in the baseline.\n"
    "LS-03: Hand-raiser vs. non-hand-raiser use cases shall be distinguished in the baseline and reflected "
    "in orchestration logic where scoring rules differ."
)

GLOSSARY_ROWS = [
    ("ATC", "Air Traffic Control — marketing data governance program for standardized activity data flows."),
    ("Tray.io", "Orchestration platform used to validate, transform, and route marketing activity data."),
    ("Snowflake", "Cloud data platform serving as the analytical landing zone for governed marketing data."),
    ("Baseline Registry", "Authoritative spreadsheet documenting fields, rules, sources, and downstream use cases."),
    ("Attribution", "Use case consuming campaign and channel tracking data for performance measurement."),
    ("Operational Reporting", "Use case consuming activity data for marketing operations dashboards and KPIs."),
    ("Lead Scoring", "Use case consuming contact and activity signals for qualification models."),
    ("Seller Enablement", "Use case providing sales teams with activity context for follow-up and engagement."),
    ("UTM", "Urchin Tracking Module — standard campaign tracking parameters (source, medium, campaign, etc.)."),
    ("MAT", "Marketing Activity Tracking — governed tracking framework for future-state parameters."),
    ("Integrate", "Content syndication partner platform for paid lead and content syndication inputs."),
    ("Hand Raiser", "Lead use case where the contact explicitly expresses purchase or engagement intent."),
    ("Non-Hand Raiser", "Lead use case where intent is implied or nurtured rather than explicitly stated."),
    ("Section 8", "Activity Detail — use-case-specific fields beyond identity, compliance, and tracking."),
]

STAKEHOLDER_ROWS = [
    ("Marketing Operations", "MO", "Owns baseline registry, change control, and activity onboarding."),
    ("Marketing Analytics / Attribution", "MA", "Defines tracking, attribution, and reporting field requirements."),
    ("Data Engineering", "DE", "Implements Tray.io orchestration and Snowflake publication pipelines."),
    ("Demand Generation", "DG", "Launches campaigns and ensures source capture aligns to baseline."),
    ("Sales Operations", "SO", "Defines seller enablement and lead scoring consumption requirements."),
    ("Compliance / Privacy", "CP", "Reviews consent, permissions, and PII-related field handling."),
    ("Program Management", "PM", "Coordinates stakeholder review, milestones, and sign-off."),
]


def set_paragraph_text(paragraph, text):
    """Replace paragraph text while preserving style."""
    if paragraph.runs:
        paragraph.runs[0].text = text
        for run in paragraph.runs[1:]:
            run.text = ""
    else:
        paragraph.add_run(text)


def replace_placeholders(doc):
    for paragraph in doc.paragraphs:
        text = paragraph.text
        if text in REPLACEMENTS:
            set_paragraph_text(paragraph, REPLACEMENTS[text])
        elif text == "Email Specific Requirements":
            set_paragraph_text(paragraph, "Online Forms Requirements")
        elif text == "Paid Social Requirements":
            set_paragraph_text(
                paragraph,
                "Attribution, Reporting, Lead Scoring and Seller Enablement Requirements",
            )
        elif text == "<An alphabetical list of terms in a particular domain of knowledge with the definitions for those terms>":
            set_paragraph_text(
                paragraph,
                "The following terms apply to the Air Traffic Control marketing data governance program.",
            )

    indices = [
        i
        for i, p in enumerate(doc.paragraphs)
        if p.text == "<List out your Specific Stakeholder Requirements>"
    ]
    if len(indices) >= 1:
        set_paragraph_text(doc.paragraphs[indices[0]], ONLINE_FORMS_REQUIREMENTS)
    if len(indices) >= 2:
        set_paragraph_text(doc.paragraphs[indices[1]], PAID_SOCIAL_REQUIREMENTS)


def populate_revision_history(table):
    rows = [
        ("2025-07-08", "0.1", "Initial draft — Air Traffic Control marketing data governance BRD", "Marketing Operations"),
        ("", "", "", ""),
        ("", "", "", ""),
        ("", "", "", ""),
        ("", "", "", ""),
    ]
    for r_idx, row_data in enumerate(rows):
        if r_idx < len(table.rows):
            for c_idx, val in enumerate(row_data):
                table.rows[r_idx].cells[c_idx].text = val


def populate_stakeholders(table):
    for r_idx, row_data in enumerate(STAKEHOLDER_ROWS):
        row_num = r_idx + 1
        if row_num < len(table.rows):
            for c_idx, val in enumerate(row_data):
                table.rows[row_num].cells[c_idx].text = val


def populate_glossary(table):
    # Keep header row, replace/add rows
    while len(table.rows) < len(GLOSSARY_ROWS) + 1:
        table.add_row()
    for r_idx, (acronym, desc) in enumerate(GLOSSARY_ROWS):
        row_num = r_idx + 1
        table.rows[row_num].cells[0].text = acronym
        table.rows[row_num].cells[1].text = desc


def main():
    doc = Document(TEMPLATE)
    replace_placeholders(doc)
    populate_revision_history(doc.tables[0])
    populate_stakeholders(doc.tables[1])
    populate_glossary(doc.tables[2])
    doc.save(OUTPUT)
    print(f"Created {OUTPUT}")


if __name__ == "__main__":
    main()
